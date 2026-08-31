import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os
import re

def clean_xml_text(text):
    if not text:
        return ""
    # Strip XML incompatible control characters
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x84\x86-\x9F]', '', text)

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def add_footer_page_number(run):
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_formatted_runs(p, text, default_font='Times New Roman', default_size=12, default_color=RGBColor(15, 23, 42), is_bold=False, is_italic=False):
    text = clean_xml_text(text)
    if not text:
        return

    tokens = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for tok in tokens:
        if not tok:
            continue
        run = p.add_run()
        run.font.name = default_font
        run.font.size = Pt(default_size)
        run.font.color.rgb = default_color
        run.bold = is_bold
        run.italic = is_italic

        if tok.startswith('***') and tok.endswith('***') and len(tok) >= 6:
            run.text = tok[3:-3]
            run.bold = True
            run.italic = True
        elif tok.startswith('**') and tok.endswith('**') and len(tok) >= 4:
            run.text = tok[2:-2]
            run.bold = True
        elif tok.startswith('*') and tok.endswith('*') and len(tok) >= 2:
            run.text = tok[1:-1]
            run.italic = True
        elif tok.startswith('`') and tok.endswith('`') and len(tok) >= 2:
            run.text = tok[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(default_size - 1)
            run.font.color.rgb = RGBColor(190, 24, 93)  # Rose/Wine code accent
        else:
            run.text = tok

def create_report_docx(filename):
    doc = Document()
    
    # Page Margins: 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Add Header & Footer
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run_left = footer_p.add_run("MERN Stack Job Portal System — Final Project Report        Page ")
        footer_run_left.font.name = 'Arial'
        footer_run_left.font.size = Pt(9)
        footer_run_left.font.color.rgb = RGBColor(100, 116, 139)
        add_footer_page_number(footer_p.add_run())

    def add_heading_styled(text, level, is_chapter=False):
        if is_chapter:
            doc.add_page_break()

        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        
        font_name = 'Arial'
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24) if not is_chapter else Pt(10)
            p.paragraph_format.space_after = Pt(14)
            add_formatted_runs(p, text, default_font=font_name, default_size=18, default_color=RGBColor(30, 58, 138), is_bold=True)
        elif level == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            add_formatted_runs(p, text, default_font=font_name, default_size=14, default_color=RGBColor(30, 58, 138), is_bold=True)
        elif level == 3:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            add_formatted_runs(p, text, default_font=font_name, default_size=12, default_color=RGBColor(51, 65, 85), is_bold=True)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_runs(p, text, default_font=font_name, default_size=11, default_color=RGBColor(71, 85, 105), is_bold=True, is_italic=True)
        return p

    def add_p(text, justify=True, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.25
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_formatted_runs(p, text, default_font='Times New Roman', default_size=12, default_color=RGBColor(15, 23, 42))
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_formatted_runs(p, text, default_font='Times New Roman', default_size=12, default_color=RGBColor(15, 23, 42))
        return p

    def add_num_list(text, num_str):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_formatted_runs(p, text, default_font='Times New Roman', default_size=12, default_color=RGBColor(15, 23, 42))
        return p

    def add_blockquote(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_formatted_runs(p, text, default_font='Times New Roman', default_size=11, default_color=RGBColor(51, 65, 85), is_italic=True)
        return p

    def add_code_block(code_lines):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        set_table_borders(tbl, color="CBD5E1", sz="6", val="single")
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        
        code_text = clean_xml_text("\n".join(code_lines))
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 41, 59)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_table_grid(headers, rows_data):
        tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl, color="CBD5E1", sz="4", val="single")
        
        hdr_cells = tbl.rows[0].cells
        for i, header_text in enumerate(headers):
            set_cell_background(hdr_cells[i], "1E3A8A")
            set_cell_margins(hdr_cells[i], top=140, bottom=140, left=140, right=140)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            add_formatted_runs(p, header_text.strip(), default_font='Arial', default_size=9.5, default_color=RGBColor(255, 255, 255), is_bold=True)

        for r_idx, row_values in enumerate(rows_data):
            row_cells = tbl.rows[r_idx + 1].cells
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_values):
                if c_idx < len(row_cells):
                    set_cell_background(row_cells[c_idx], bg_color)
                    set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=140, right=140)
                    p = row_cells[c_idx].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_after = Pt(0)
                    add_formatted_runs(p, str(val).strip(), default_font='Times New Roman', default_size=9.5, default_color=RGBColor(15, 23, 42))

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Read markdown and build docx elements
    md_file = r"c:\Users\HomePC\Desktop\job-portal\job-portal\Final_Project_Report.md"
    if os.path.exists(md_file):
        with open(md_file, "r", encoding="utf-8") as f:
            lines = f.readlines()
        
        in_table = False
        in_code = False
        code_lines = []
        table_headers = []
        table_rows = []

        for line in lines:
            line_str = line.rstrip("\r\n")
            stripped = line_str.strip()

            # Handle Code Blocks
            if stripped.startswith("```"):
                if in_code:
                    add_code_block(code_lines)
                    in_code = False
                    code_lines = []
                else:
                    if in_table and table_headers:
                        add_table_grid(table_headers, table_rows)
                        in_table = False
                        table_headers = []
                        table_rows = []
                    in_code = True
                    code_lines = []
                continue

            if in_code:
                code_lines.append(line_str)
                continue

            # Empty lines
            if not stripped:
                if in_table and table_headers:
                    add_table_grid(table_headers, table_rows)
                    in_table = False
                    table_headers = []
                    table_rows = []
                continue

            # Handle Tables
            if stripped.startswith("|"):
                parts = [p.strip() for p in stripped.split("|")[1:-1]]
                if not parts or all(set(p) <= set("-: ") for p in parts):
                    continue  # Separator line
                if not in_table:
                    in_table = True
                    table_headers = parts
                    table_rows = []
                else:
                    table_rows.append(parts)
                continue
            elif in_table:
                add_table_grid(table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []

            # Handle Markdown Headings & Elements
            if stripped.startswith("# "):
                heading_title = stripped[2:].strip()
                is_chapter = any(kw in heading_title.upper() for kw in ["CHAPTER", "PART C", "APPENDIX"])
                add_heading_styled(heading_title, 1, is_chapter=is_chapter)
            elif stripped.startswith("## "):
                add_heading_styled(stripped[3:].strip(), 2)
            elif stripped.startswith("### "):
                add_heading_styled(stripped[4:].strip(), 3)
            elif stripped.startswith("#### "):
                add_heading_styled(stripped[5:].strip(), 4)
            elif stripped.startswith("> "):
                add_blockquote(stripped[2:].strip())
            elif stripped.startswith("- ") or stripped.startswith("* "):
                add_bullet(stripped[2:].strip())
            elif re.match(r'^\d+\.\s', stripped):
                match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
                if match:
                    add_num_list(match.group(2), match.group(1))
                else:
                    add_bullet(stripped)
            elif stripped == "---":
                continue
            else:
                add_p(stripped)

        if in_table and table_headers:
            add_table_grid(table_headers, table_rows)
        if in_code and code_lines:
            add_code_block(code_lines)

    doc.save(filename)
    print(f"Successfully generated Word document: {filename}")

if __name__ == "__main__":
    out_docx = r"c:\Users\HomePC\Desktop\job-portal\job-portal\Final_Project_Report.docx"
    create_report_docx(out_docx)
